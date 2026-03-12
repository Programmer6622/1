import re
import shutil
import subprocess
import tempfile
import zipfile
from collections import defaultdict
from pathlib import Path

from .errors import ConversionError


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".txt":
        return _read_text(path)
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".doc":
        return _extract_doc(path)
    if ext in {".xlsx", ".xlsm", ".xls", ".csv"}:
        raise ConversionError("Spreadsheet files are not supported.")
    if ext in {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}:
        raise ConversionError("Images are not supported.")
    raise ConversionError(f"Unsupported file type: {ext}")


def _read_text(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except Exception as exc:
        raise ConversionError("pdfplumber is required to read PDF files.") from exc
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text:
                parts.append(text)
    return "\n".join(parts)


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception:
        return _extract_docx_zip(path)
    doc = Document(str(path))
    parts = []
    list_counters: dict[tuple[int, int], int] = defaultdict(int)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        num_info = _get_docx_num_info(para)
        if num_info is not None:
            num_id, ilvl = num_info
            key = (num_id, ilvl)
            list_counters[key] += 1
            idx = list_counters[key]
            if _looks_like_question_sentence(text):
                text = f"{idx}. {text}"
            else:
                text = f"{idx}) {text}"
        is_list_option = num_info is not None and not _looks_like_question_sentence(text)
        has_correct_style = _paragraph_has_red_run(para) or _paragraph_has_highlight_run(para)
        if has_correct_style and (is_list_option or _looks_like_option_line(text)):
            text = f"{text} *"
        parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells if cell.text]
            if cells:
                parts.append(" ".join(cells))
    return "\n".join(parts)


def _looks_like_option_line(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if re.fullmatch(r"[\d\W_]+", t):
        return False
    if t.endswith(("?", ":")):
        return False
    # Typical answer options are relatively short one-line statements.
    return len(t.split()) <= 18


def _paragraph_has_red_run(para) -> bool:
    for run in para.runs:
        color = run.font.color
        rgb = getattr(color, "rgb", None) if color is not None else None
        if rgb is None:
            continue
        rgb_s = str(rgb).upper()
        if rgb_s in {"FF0000", "C00000"}:
            return True
    return False


def _paragraph_has_highlight_run(para) -> bool:
    for run in para.runs:
        if run.font.highlight_color is not None:
            return True
        rpr = run._r.rPr
        if rpr is not None and rpr.highlight is not None:
            return True
    return False


def _get_docx_num_info(para) -> tuple[int, int] | None:
    p_pr = para._p.pPr
    if p_pr is None or p_pr.numPr is None:
        return None
    num_pr = p_pr.numPr
    if num_pr.numId is None:
        return None
    num_id = int(num_pr.numId.val)
    ilvl = int(num_pr.ilvl.val) if num_pr.ilvl is not None else 0
    return num_id, ilvl


def _looks_like_question_sentence(text: str) -> bool:
    t = text.strip()
    if t.endswith(("?", ":")):
        return True
    return bool(
        re.match(
            r"(?i)^(определите|укажите|выберите|назовите|какой|какая|какие|кто|что|когда|сколько)\b",
            t,
        )
    )


def _extract_docx_zip(path: Path) -> str:
    with zipfile.ZipFile(path) as zf:
        data = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    data = data.replace("</w:p>", "\n").replace("</w:tr>", "\n")
    data = data.replace("</w:tab>", "\t")
    text = re.sub(r"<[^>]+>", "", data)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _extract_doc(path: Path) -> str:
    if shutil.which("soffice"):
        with tempfile.TemporaryDirectory() as tmpdir:
            cmd = [
                "soffice",
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                tmpdir,
                str(path),
            ]
            subprocess.run(cmd, check=True, capture_output=True)
            txt_path = Path(tmpdir) / f"{path.stem}.txt"
            if not txt_path.exists():
                raise ConversionError("LibreOffice did not produce a text file.")
            return _read_text(txt_path)
    if shutil.which("antiword"):
        result = subprocess.run(
            ["antiword", str(path)],
            check=True,
            capture_output=True,
            text=True,
        )
        return result.stdout
    raise ConversionError("DOC files require LibreOffice (soffice) or antiword.")

